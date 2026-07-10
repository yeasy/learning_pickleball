#!/usr/bin/env python3
"""Build Chinese and English DOCX editions from SUMMARY-ordered Markdown."""

from __future__ import annotations

import argparse
import os
import re
import subprocess
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path

try:
    from PIL import Image
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm
except ImportError as exc:  # pragma: no cover - exercised by the CLI environment
    raise SystemExit(
        "DOCX dependencies are missing. Run: python3 -m pip install -r requirements-docx.txt"
    ) from exc

ROOT = Path(__file__).resolve().parents[1]
DOCX_TOOLS = Path(__file__).resolve().parent / "docx"
sys.path.insert(0, str(DOCX_TOOLS))
from build_reference_doc import create_reference_doc  # noqa: E402


SUMMARY_RE = re.compile(r"^\s*[-*]\s+\[[^]]+\]\(([^)#?]+\.md)(?:#[^)]*)?\)")
REMOTE_BADGE_RE = re.compile(r"\[!\[[^]]*\]\(https?://[^)]*\)\]\([^)]*\)")
REMOTE_IMAGE_RE = re.compile(r"!\[[^]]*\]\(https?://[^)]*\)")
LANGUAGE_SWITCH_RE = re.compile(r"\A\[(?:English Version|中文版)\]\([^)]+\)\s*")
H1_RE = re.compile(r"^#\s+(.+?)\s*$", re.MULTILINE)


class BuildError(RuntimeError):
    """A user-actionable DOCX build failure."""


def summary_files(language: str) -> list[Path]:
    language_root = ROOT / language
    summary = language_root / "SUMMARY.md"
    paths = []
    for line in summary.read_text(encoding="utf-8").splitlines():
        match = SUMMARY_RE.match(line)
        if match:
            paths.append(language_root / match.group(1))
    if len(paths) != 28:
        raise BuildError(f"{language}/SUMMARY.md must contain 28 Markdown entries; found {len(paths)}")
    if len(paths) != len(set(paths)):
        raise BuildError(f"{language}/SUMMARY.md contains duplicate entries")
    missing = [str(path.relative_to(ROOT)) for path in paths if not path.is_file()]
    if missing:
        raise BuildError(f"missing SUMMARY inputs: {missing}")
    return paths


def _local_image_target(source: Path, raw_target: str) -> Path | None:
    target = raw_target.strip().split("#", 1)[0].split("?", 1)[0]
    if target.startswith(("http://", "https://", "data:", "/")):
        return None
    path = (source.parent / target).resolve()
    try:
        path.relative_to(ROOT)
    except ValueError as exc:
        raise BuildError(f"image escapes repository: {source.relative_to(ROOT)} -> {raw_target}") from exc
    if not path.is_file():
        raise BuildError(f"missing image: {source.relative_to(ROOT)} -> {raw_target}")
    return path


def normalize_markdown(source: Path) -> str:
    text = source.read_text(encoding="utf-8")
    text = LANGUAGE_SWITCH_RE.sub("", text, count=1)
    text = REMOTE_BADGE_RE.sub("", text)
    text = REMOTE_IMAGE_RE.sub("", text)

    def replace_markdown_image(match: re.Match[str]) -> str:
        alt, target = match.group(1), match.group(2)
        local = _local_image_target(source, target)
        return match.group(0) if local is None else f"![{alt}]({local.as_posix()})"

    def replace_html_image(match: re.Match[str]) -> str:
        target = match.group(1)
        local = _local_image_target(source, target)
        return match.group(0) if local is None else match.group(0).replace(target, local.as_posix())

    text = re.sub(r"!\[([^]]*)\]\(([^)]+)\)", replace_markdown_image, text)
    text = re.sub(r'<img\s+[^>]*src="([^"]+)"[^>]*>', replace_html_image, text)
    return text.strip() + "\n"


def crop_cover(language: str, output: Path) -> None:
    source = ROOT / "_images/cover.jpg"
    with Image.open(source) as image:
        midpoint = image.width // 2
        box = (0, 0, midpoint, image.height) if language == "cn" else (
            midpoint,
            0,
            image.width,
            image.height,
        )
        cover = image.crop(box).convert("RGB")
        cover.save(output, format="PNG", optimize=False, compress_level=9)


def _fixed_datetime() -> datetime:
    raw = os.environ.get("SOURCE_DATE_EPOCH")
    if raw and raw.isdigit() and int(raw) >= 315532800:
        return datetime.fromtimestamp(int(raw), timezone.utc).replace(tzinfo=None)
    return datetime(2000, 1, 1, 0, 0, 0)


def _enable_field_updates(document: Document) -> None:
    settings = document.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def _normalize_bullet_glyphs(document: Document) -> None:
    numbering = document.part.numbering_part.element
    for level in numbering.findall(f".//{qn('w:lvl')}"):
        number_format = level.find(qn("w:numFmt"))
        level_text = level.find(qn("w:lvlText"))
        if (
            number_format is None
            or number_format.get(qn("w:val")) != "bullet"
            or level_text is None
        ):
            continue
        level_text.set(qn("w:val"), "•")
        run_properties = level.find(qn("w:rPr"))
        if run_properties is not None:
            level.remove(run_properties)


def _materialize_toc(document: Document, language: str, titles: list[str]) -> None:
    body = document._element.body
    toc = next(
        (
            child
            for child in body
            if child.tag == qn("w:sdt")
            and child.find(f".//{qn('w:docPartGallery')}") is not None
            and child.find(f".//{qn('w:docPartGallery')}").get(qn("w:val"))
            == "Table of Contents"
        ),
        None,
    )
    if toc is None:
        raise BuildError("pandoc output is missing the table-of-contents container")
    content = toc.find(qn("w:sdtContent"))
    if content is None:
        raise BuildError("pandoc table of contents has no content container")
    for child in list(content):
        content.remove(child)

    heading = document.add_paragraph(
        "目录" if language == "cn" else "Table of Contents",
        style="TOC Heading",
    )
    content.append(heading._p)
    for title in titles:
        entry = document.add_paragraph(title, style="TOC 1")
        content.append(entry._p)


def add_front_matter(
    body_docx: Path,
    cover: Path,
    language: str,
    toc_titles: list[str],
    output: Path,
) -> None:
    document = Document(body_docx)
    _normalize_bullet_glyphs(document)
    _materialize_toc(document, language, toc_titles)
    body = document._element.body
    anchor = body[0]
    title = "学打匹克球" if language == "cn" else "Learning Pickleball"
    source_note = (
        "内容由本仓库 Markdown 按 SUMMARY.md 顺序生成。"
        if language == "cn"
        else "Generated from this repository's Markdown in SUMMARY.md order."
    )

    front = []
    cover_paragraph = document.add_paragraph()
    cover_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_paragraph.add_run().add_picture(str(cover), width=Mm(150))
    front.append(cover_paragraph._p)
    page_break = document.add_paragraph()
    page_break.add_run().add_break(WD_BREAK.PAGE)
    front.append(page_break._p)
    title_paragraph = document.add_paragraph(title, style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.append(title_paragraph._p)
    author = document.add_paragraph("yeasy")
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.append(author._p)
    note = document.add_paragraph(source_note)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.append(note._p)
    second_break = document.add_paragraph()
    second_break.add_run().add_break(WD_BREAK.PAGE)
    front.append(second_break._p)
    for element in front:
        anchor.addprevious(element)

    fixed = _fixed_datetime()
    properties = document.core_properties
    properties.title = title
    properties.author = "yeasy"
    properties.subject = "Pickleball instruction"
    properties.created = fixed
    properties.modified = fixed
    _enable_field_updates(document)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def build_language(language: str, output_dir: Path, pandoc: str) -> Path:
    sources = summary_files(language)
    toc_titles = []
    for source in sources:
        match = H1_RE.search(source.read_text(encoding="utf-8"))
        if not match:
            raise BuildError(f"missing H1 in {source.relative_to(ROOT)}")
        toc_titles.append(match.group(1))
    output = output_dir / f"learning_pickleball-{language}.docx"
    with tempfile.TemporaryDirectory(prefix=f"learning-pickleball-{language}-") as tmp:
        temp = Path(tmp)
        combined = temp / "book.md"
        combined.write_text("\n\n".join(normalize_markdown(path) for path in sources), encoding="utf-8")
        reference = temp / "reference.docx"
        create_reference_doc(language, reference)
        cover = temp / f"cover-{language}.png"
        crop_cover(language, cover)
        body_docx = temp / "body.docx"
        command = [
            pandoc,
            str(combined),
            "--from=gfm",
            "--to=docx",
            "--toc",
            "--toc-depth=3",
            f"--reference-doc={reference}",
            f"--resource-path={ROOT / language}{os.pathsep}{ROOT}",
            f"--output={body_docx}",
        ]
        result = subprocess.run(command, cwd=ROOT, text=True, capture_output=True, check=False)
        if result.returncode:
            raise BuildError(f"pandoc failed for {language}: {result.stderr.strip()}")
        staged = temp / output.name
        add_front_matter(body_docx, cover, language, toc_titles, staged)
        output_dir.mkdir(parents=True, exist_ok=True)
        os.replace(staged, output)
    if not output.is_file() or output.stat().st_size == 0:
        raise BuildError(f"empty output: {output}")
    print(f"built {output}")
    return output


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--lang", choices=("cn", "en"), action="append")
    parser.add_argument("--output-dir", type=Path, default=ROOT)
    parser.add_argument("--pandoc", default="pandoc")
    args = parser.parse_args()
    languages = args.lang or ["cn", "en"]
    try:
        for language in languages:
            build_language(language, args.output_dir.resolve(), args.pandoc)
    except (BuildError, OSError, ValueError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
